"""
resume_source.py — Extracts candidate data from a resume file (PDF or DOCX).

Supported formats: .pdf, .docx
Unsupported/corrupt files degrade to [] (no profile) rather than crashing.
"""

from __future__ import annotations
import logging
import re
from pathlib import Path
from typing import Optional

from transformer.schema import (
    CandidateProfile, LinksEntry, SkillEntry, ExperienceEntry,
    EducationEntry, ProvenanceEntry,
)
from transformer.normalizers import (
    normalize_name, normalize_email, normalize_phone,
    normalize_skill, normalize_date,
)

logger = logging.getLogger(__name__)
SOURCE_NAME = "resume"

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(\+?\d[\d\s\-().]{7,}\d)")
_LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_%]+/?", re.IGNORECASE)
_GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9\-]+/?", re.IGNORECASE)
_PORTFOLIO_RE = re.compile(r"(?:https?://)?(?:www\.)?[A-Za-z0-9\-]+\.(?:dev|me|io|com)(?:/[^\s,]*)?", re.IGNORECASE)

# Section headers we look for, generically — any resume that uses
# reasonably standard section names will match. Order matters for
# splitting but not for which sections exist.
_SECTION_PATTERNS: dict[str, re.Pattern] = {
    "experience": re.compile(
        r"^\s*(work\s+experience|professional\s+experience|experience|employment\s+history)\s*:?\s*$",
        re.IGNORECASE,
    ),
    "education": re.compile(r"^\s*(education|academic\s+background)\s*:?\s*$", re.IGNORECASE),
    "skills": re.compile(
        r"^\s*(skills|technical\s+skills|technologies|core\s+competencies)\s*:?\s*$",
        re.IGNORECASE,
    ),
    "summary": re.compile(
        r"^\s*(summary|profile|about|objective|professional\s+summary)\s*:?\s*$",
        re.IGNORECASE,
    ),
}


# ─── File reading ────────────────────────────────────────────────────────────

def _read_pdf_text(path: Path) -> Optional[str]:
    try:
        import pdfplumber
    except ImportError:
        logger.error("[resume] pdfplumber not installed; cannot read PDF resumes")
        return None
    try:
        text_parts: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.error(f"[resume] Failed to read PDF {path}: {e}")
        return None


def _read_docx_text(path: Path) -> Optional[str]:
    try:
        import docx
    except ImportError:
        logger.error("[resume] python-docx not installed; cannot read DOCX resumes")
        return None
    try:
        document = docx.Document(str(path))
        lines: list[str] = []
        for para in document.paragraphs:
            lines.append(para.text)
        # Tables sometimes hold contact info / skills in modern resume templates
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text)
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"[resume] Failed to read DOCX {path}: {e}")
        return None


def _load_text(path: Path) -> Optional[str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf_text(path)
    if suffix == ".docx":
        return _read_docx_text(path)
    logger.warning(f"[resume] Unsupported resume file type: {suffix}")
    return None


# ─── Generic section splitting ───────────────────────────────────────────────

def _split_sections(text: str) -> dict[str, str]:
    """
    Walk lines; whenever a line matches a known section header pattern,
    start a new section. Everything before the first recognized header is
    the 'header' block (name, contact info, sometimes a one-line title).
    """
    lines = text.splitlines()
    sections: dict[str, list[str]] = {"header": []}
    current = "header"
    for line in lines:
        matched = None
        for name, pattern in _SECTION_PATTERNS.items():
            if pattern.match(line.strip()):
                matched = name
                break
        if matched:
            current = matched
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {k: "\n".join(v).strip() for k, v in sections.items()}


# ─── Header parsing (name, contact, links) ──────────────────────────────────

def _parse_contact_info(full_text: str) -> dict[str, Optional[str]]:
    """Scan the WHOLE document (not just header) for contact patterns —
    resumes place these inconsistently (header line, footer, sidebar)."""
    info: dict[str, Optional[str]] = {
        "email": None, "phone": None, "linkedin": None,
        "github": None, "portfolio": None,
    }

    email_m = _EMAIL_RE.search(full_text)
    if email_m:
        info["email"] = email_m.group(0)

    linkedin_m = _LINKEDIN_RE.search(full_text)
    if linkedin_m:
        url = linkedin_m.group(0)
        info["linkedin"] = url if url.startswith("http") else f"https://{url}"

    github_m = _GITHUB_RE.search(full_text)
    if github_m:
        url = github_m.group(0)
        info["github"] = url if url.startswith("http") else f"https://{url}"

    # Phone: search line-by-line near the top to avoid false-positives from
    # dates like "2019 - 2023" later in the document (those rarely look
    # like phone numbers due to the digit-count/format check in normalize_phone,
    # but we still bias toward the first 15 lines for precision).
    top_lines = "\n".join(full_text.splitlines()[:15])
    phone_m = _PHONE_RE.search(top_lines) or _PHONE_RE.search(full_text)
    if phone_m:
        info["phone"] = phone_m.group(0)

    # Portfolio: any other URL-looking token that isn't linkedin/github
    for m in _PORTFOLIO_RE.finditer(full_text):
        candidate = m.group(0)
        if "linkedin.com" in candidate.lower() or "github.com" in candidate.lower():
            continue
        info["portfolio"] = candidate if candidate.startswith("http") else f"https://{candidate}"
        break

    return info


def _parse_name(header_block: str) -> Optional[str]:
    """
    Generic heuristic: the first non-empty line in the header block that
    is NOT itself an email/phone/url is treated as the candidate's name,
    provided it looks name-like (2-5 words, mostly alphabetic, no digits).
    """
    for line in header_block.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if _EMAIL_RE.search(stripped) or _LINKEDIN_RE.search(stripped) or _GITHUB_RE.search(stripped):
            continue
        if any(ch.isdigit() for ch in stripped):
            continue
        words = stripped.split()
        if 1 <= len(words) <= 5 and all(w.replace(".", "").replace(",", "").isalpha() for w in words):
            return stripped
    return None


def _parse_headline(header_block: str, summary_block: str, name: Optional[str]) -> Optional[str]:
    """
    Headline candidates, in priority order:
      1. The line right after the name in the header (a common resume
         pattern: "Jane Doe \n Senior Backend Engineer").
      2. The first sentence of a Summary/About/Objective section.
    """
    lines = [ln.strip() for ln in header_block.splitlines() if ln.strip()]
    if name and lines:
        try:
            name_idx = lines.index(name)
            if name_idx + 1 < len(lines):
                candidate = lines[name_idx + 1]
                if not _EMAIL_RE.search(candidate) and not _PHONE_RE.search(candidate):
                    return candidate
        except ValueError:
            pass

    if summary_block:
        first_sentence = re.split(r"(?<=[.!?])\s+", summary_block.strip())[0]
        if first_sentence:
            return first_sentence[:300]

    return None


# ─── Experience / Education / Skills parsing (shared patterns with linkedin-style parsing) ──

_EXP_TITLE_AT_RE = re.compile(r"^(.*?)\s*[-–—|]\s*(.+)$")
_EXP_TITLE_AT_WORD_RE = re.compile(r"^(.*?)\s+(?:at|@)\s+(.+)$", re.IGNORECASE)
_DATE_RANGE_RE = re.compile(
    r"([A-Za-z]+\s+\d{4}|\d{4})\s*[-–—to]+\s*(Present|present|Current|current|[A-Za-z]+\s+\d{4}|\d{4})"
)


def _parse_experience(block: str) -> list[ExperienceEntry]:
    """
    Generic, format-tolerant parser: splits on blank lines into entries.
    First line of an entry is treated as "Title at Company" or
    "Title - Company" or "Title | Company"; a following line containing a
    recognizable date range becomes start/end; remaining lines become the
    summary. Entries that yield neither title nor company are skipped.
    """
    if not block:
        return []
    entries: list[ExperienceEntry] = []
    chunks = re.split(r"\n\s*\n", block.strip())
    for chunk in chunks:
        lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
        if not lines:
            continue

        title_line = lines[0]
        m = _EXP_TITLE_AT_WORD_RE.match(title_line) or _EXP_TITLE_AT_RE.match(title_line)
        if m:
            title, company = m.group(1).strip(), m.group(2).strip()
        else:
            title, company = title_line, ""

        # Strip a trailing date range that sometimes rides on the title line itself
        date_on_title = _DATE_RANGE_RE.search(company) or _DATE_RANGE_RE.search(title)

        start, end = None, None
        summary_lines: list[str] = []
        remaining_lines = lines[1:]
        for ln in remaining_lines:
            dm = _DATE_RANGE_RE.search(ln)
            if dm and start is None:
                start = normalize_date(dm.group(1))
                end = normalize_date(dm.group(2))
            else:
                summary_lines.append(ln)

        if date_on_title and start is None:
            start = normalize_date(date_on_title.group(1))
            end = normalize_date(date_on_title.group(2))
            # Remove the date fragment from company/title text
            company = _DATE_RANGE_RE.sub("", company).strip(" ,-")
            title = _DATE_RANGE_RE.sub("", title).strip(" ,-")

        if not title and not company:
            continue

        entries.append(ExperienceEntry(
            company=company,
            title=title,
            start=start,
            end=end,
            summary=" ".join(summary_lines) if summary_lines else None,
        ))
    return entries


_EDU_YEAR_RE = re.compile(r"(\d{4})")


def _parse_education(block: str) -> list[EducationEntry]:
    """
    Generic parser: blank-line separated entries. First line = institution.
    Optional second line = "Degree, Field, Year" or "Degree in Field, Year"
    or any comma-separated mix; we extract the year if present and treat
    the remaining comma-separated parts as degree/field in order.
    """
    if not block:
        return []
    entries: list[EducationEntry] = []
    chunks = re.split(r"\n\s*\n", block.strip())
    for chunk in chunks:
        lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
        if not lines:
            continue
        institution = lines[0]
        degree, field, end_year = None, None, None

        detail_lines = lines[1:]
        for detail in detail_lines:
            year_m = _EDU_YEAR_RE.search(detail)
            cleaned = detail
            if year_m:
                end_year = int(year_m.group(1))
                cleaned = _EDU_YEAR_RE.sub("", detail).strip(" ,()-")

            cleaned = re.sub(r"\bin\b", ",", cleaned, flags=re.IGNORECASE)
            parts = [p.strip() for p in cleaned.split(",") if p.strip()]
            if parts and degree is None:
                degree = parts[0]
            if len(parts) >= 2 and field is None:
                field = parts[1]

        entries.append(EducationEntry(
            institution=institution, degree=degree, field=field, end_year=end_year,
        ))
    return entries


def _parse_skills(block: str) -> list[str]:
    """Skills section: comma/pipe/semicolon separated, or one-per-line, or
    bullet-prefixed lines — handles whichever format the resume used."""
    if not block:
        return []
    cleaned_lines = [re.sub(r"^[•\-*▪◦]\s*", "", ln.strip()) for ln in block.splitlines() if ln.strip()]
    joined = " ".join(cleaned_lines)
    for sep in [",", ";", "|", "•"]:
        if joined.count(sep) >= 1:
            return [s.strip() for s in joined.split(sep) if s.strip()]
    return cleaned_lines


# ─── Main extractor ──────────────────────────────────────────────────────────

def extract(source_path: str) -> list[CandidateProfile]:
    """
    Parse a resume file (PDF or DOCX) into a single-element list containing
    a partial CandidateProfile, or [] if the file is missing, unsupported,
    unreadable, or contains no extractable content. Never raises.
    """
    path = Path(source_path)
    if not path.exists():
        logger.warning(f"[resume] File not found: {source_path}")
        return []

    text = _load_text(path)
    if not text or not text.strip():
        logger.warning(f"[resume] No extractable text in {source_path}")
        return []

    profile = _build_profile(text, source_path)
    if profile is None:
        return []

    logger.info(f"[resume] Extracted profile from {source_path}")
    return [profile]


def _build_profile(text: str, source_path: str) -> Optional[CandidateProfile]:
    sections = _split_sections(text)
    header_block = sections.get("header", "")

    provenance: list[ProvenanceEntry] = []

    def record(field: str, method: str = "direct") -> None:
        provenance.append(ProvenanceEntry(field=field, source=SOURCE_NAME, method=method))

    # Name
    raw_name = _parse_name(header_block)
    full_name: Optional[str] = None
    if raw_name:
        full_name = normalize_name(raw_name)
        if full_name:
            record("full_name")

    # Contact info (scanned across whole doc — resumes place this inconsistently)
    contact = _parse_contact_info(text)

    emails: list[str] = []
    if contact["email"]:
        em = normalize_email(contact["email"])
        if em:
            emails.append(em)
            record("emails")

    phones: list[str] = []
    if contact["phone"]:
        ph = normalize_phone(contact["phone"])
        if ph:
            phones.append(ph)
            record("phones", "normalized")

    links: Optional[LinksEntry] = None
    if contact["linkedin"] or contact["github"] or contact["portfolio"]:
        links = LinksEntry(
            linkedin=contact["linkedin"],
            github=contact["github"],
            portfolio=contact["portfolio"],
        )
        record("links")

    # Headline
    headline = _parse_headline(header_block, sections.get("summary", ""), raw_name)
    if headline:
        record("headline")

    # Experience
    experience = _parse_experience(sections.get("experience", ""))
    if experience:
        record("experience", "normalized")

    # Education
    education = _parse_education(sections.get("education", ""))
    if education:
        record("education", "normalized")

    # Skills
    raw_skills = _parse_skills(sections.get("skills", ""))
    skills: list[SkillEntry] = []
    for rs in raw_skills:
        canonical = normalize_skill(rs)
        if canonical:
            skills.append(SkillEntry(name=canonical, confidence=0.8, sources=[SOURCE_NAME]))
    if skills:
        record("skills", "normalized")

    # years_experience: inferred from earliest experience start year, never
    # stated directly by a resume, so we tag it as inferred and keep it
    # conservative (current year - earliest start year).
    years_experience: Optional[float] = None
    starts = [int(e.start[:4]) for e in experience if e.start and e.start[:4].isdigit()]
    if starts:
        from datetime import date
        years_experience = float(date.today().year - min(starts))
        record("years_experience", "inferred")

    if not any([full_name, emails, phones, headline, experience, education, skills]):
        return None  # Nothing usable — degrade to no profile, never fabricate

    cid = f"resume_{Path(source_path).stem}"

    filled = sum([
        bool(full_name), bool(emails), bool(phones),
        bool(headline), bool(experience), bool(education), bool(skills),
    ])
    confidence = min(1.0, filled / 7.0)

    return CandidateProfile(
        candidate_id=cid,
        full_name=full_name,
        emails=emails,
        phones=phones,
        location=None,  # Resumes rarely state structured location reliably; left for other sources to supply
        links=links,
        headline=headline,
        years_experience=years_experience,
        skills=skills,
        experience=experience,
        education=education,
        provenance=provenance,
        overall_confidence=confidence,
    )