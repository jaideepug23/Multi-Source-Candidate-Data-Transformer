from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

p = doc.add_paragraph()
r = p.add_run("Aarav Sharma")
r.bold = True
r.font.size = Pt(16)

doc.add_paragraph("Mechanical Design Engineer")

doc.add_heading("Summary", level=2)
doc.add_paragraph(
    "Mechanical engineer with experience in automotive component design and CAD modeling for production vehicle programs."
)

doc.add_heading("Experience", level=2)

doc.add_paragraph("Senior Design Engineer at Tata Motors")
doc.add_paragraph("Apr 2021 - Present")
doc.add_paragraph(
    "Leading chassis component design for the commercial vehicle division, working closely with manufacturing on tolerancing and DFM reviews."
)

doc.add_paragraph()

doc.add_paragraph("Graduate Engineer Trainee at Tata Motors")
doc.add_paragraph("Jun 2019 - Mar 2021")
doc.add_paragraph(
    "Rotated across powertrain and chassis teams; contributed to suspension geometry optimization."
)

doc.add_heading("Education", level=2)

doc.add_paragraph("Pune Institute of Engineering and Technology")
doc.add_paragraph("B.E., Mechanical Engineering, 2019")

doc.add_heading("Skills", level=2)

doc.add_paragraph(
    "CATIA, SolidWorks, GD&T, DFMEA, Automotive Design"
)

doc.save("resume_aarav_sharma_namesake.docx")

print("Written resume_aarav_sharma_namesake.docx")