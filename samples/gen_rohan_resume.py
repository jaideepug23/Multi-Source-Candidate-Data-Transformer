from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# Header
p = doc.add_paragraph()
r = p.add_run("Rohan Verma")
r.bold = True
r.font.size = Pt(16)

doc.add_paragraph("Machine Learning Research Engineer | NLP")
doc.add_paragraph("rohan.verma@sric.iitd.ac.in | +91 96500 12345")
doc.add_paragraph("linkedin.com/in/rohanverma-iitd")

doc.add_paragraph()

# Summary
h = doc.add_heading("Summary", level=2)
doc.add_paragraph(
    "Machine learning researcher with a focus on natural language processing "
    "for low-resource Indian languages, currently pursuing graduate research "
    "at IIT Delhi."
)

# Experience
doc.add_heading("Experience", level=2)

doc.add_paragraph("Research Assistant at Indian Institute of Technology Delhi")
doc.add_paragraph("Jul 2023 - Present")
doc.add_paragraph(
    "Leading research on low-resource NLP for Hindi and Marathi under the "
    "Speech and NLP Lab, with two accepted publications."
)

doc.add_paragraph()

doc.add_paragraph("Machine Learning Intern at Samsung R&D Institute India")
doc.add_paragraph("Jan 2022 - Jun 2022")
doc.add_paragraph(
    "Built and evaluated transformer-based intent classification models for "
    "an on-device voice assistant."
)

# Education
doc.add_heading("Education", level=2)

doc.add_paragraph("Indian Institute of Technology Delhi")
doc.add_paragraph("M.Tech, Computer Science and Engineering, 2025")

doc.add_paragraph()

doc.add_paragraph("Indian Institute of Technology Delhi")
doc.add_paragraph("B.Tech, Computer Science and Engineering, 2023")

# Skills
doc.add_heading("Skills", level=2)

doc.add_paragraph(
    "Python, Machine Learning, PyTorch, NLP, TensorFlow, Pandas, SQL"
)

# Save
doc.save("resume_rohan_verma.docx")

print("Written resume_rohan_verma.docx")